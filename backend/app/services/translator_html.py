"""
translator_html.py  —— 改进版

目标：
- 中心策略：HTML 中间格式 + 原始 drawing/media overlay
- 修复/改进：表格合并单元格(gridSpan/vMerge)、浮动图片drawing回写、media回拷、AI注释剥离
- DeepSeek 调用：超时 120s、长文本分块、严格指令要求保持标签、不生成说明

注意：DOCX 的 DrawingML / SmartArt / VML 非常复杂；本脚本采取“尽力恢复原始 XML 片段并 overlay media/rels”的方式，
这在绝大多数文档能显著提升保真度，但并不保证 100% 在所有复杂自定义文档中完全一致。
"""

import os
import re
import io
import json
import zipfile
import shutil
import tempfile
from pathlib import Path
from typing import List, Dict, Any

import requests
from dotenv import load_dotenv
from bs4 import BeautifulSoup, NavigableString
from lxml import etree
from PIL import Image

from docx import Document
from docx.shared import RGBColor, Inches

# ----------------------------
# 配置与常量
# ----------------------------
THIS_FILE = Path(__file__).resolve()
PROJECT_ROOT = THIS_FILE.parents[2] if len(THIS_FILE.parents) >= 3 else Path.cwd()
dotenv_path = PROJECT_ROOT / ".env"
if dotenv_path.exists():
    load_dotenv(dotenv_path)
else:
    load_dotenv()

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
DEEPSEEK_API_URL = os.getenv("DEEPSEEK_API_URL", "https://api.deepseek.com/v1/chat/completions")

WP_NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'v': 'urn:schemas-microsoft-com:vml',
    'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
}

EMU_PER_INCH = 914400  # EMU -> inch (Word unit)

# ----------------------------
# Utilities
# ----------------------------
def emu_to_inches(emu_val):
    try:
        return float(emu_val) / EMU_PER_INCH
    except Exception:
        return None

def clean_ai_comments(html: str) -> str:
    """
    去掉 AI 可能插入的说明性段落，例如:
    - Note: ...
    - （注：...）
    - Here is the translated ...
    返回清洗后的 HTML 字符串（尽量保留标签）
    """
    if not html:
        return html
    soup = BeautifulSoup(html, "html.parser")
    bad_start_re = re.compile(r'^\s*(Note[:：]|（注：|Here is|Since the original|The HTML content|The original)', flags=re.I)
    # remove block tags that look like commentary
    for tag in list(soup.find_all(['p','div'])):
        txt = tag.get_text(separator=' ', strip=True)
        if not txt:
            tag.decompose()
            continue
        if bad_start_re.search(txt):
            tag.decompose()
    # remove stray lines
    lines = str(soup).splitlines()
    filtered = [l for l in lines if not bad_start_re.search(l.strip())]
    return "\n".join(filtered).strip()

def strip_empty_tags(html: str) -> str:
    """
    删除空的 p/div/span，保留含文本或含 img 的标签
    """
    if not html:
        return html
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup.find_all():
        if tag.name in ("p","div","span","section","article","li"):
            if not tag.get_text(strip=True) and not tag.find('img'):
                tag.decompose()
    return str(soup)

# ----------------------------
# DeepSeek wrapper（支持分块 & 超时 & 强制保留标签）
# ----------------------------
def deepseek_translate_html(html_fragments: List[str], src_lang="zh", tgt_lang="ja") -> List[str]:
    """
    html_fragments: 每项是一个 HTML 片段（含标签）
    返回: 对应的翻译后 HTML（已做注释清洗）
    """
    if not DEEPSEEK_API_KEY:
        print("⚠️ DEEPSEEK_API_KEY 未配置，使用 MOCK 翻译。")
        return [clean_ai_comments(strip_empty_tags(f"{frag}")) for frag in html_fragments]

    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {DEEPSEEK_API_KEY}"}
    results = []
    for frag in html_fragments:
        # skip empty
        if not frag or not BeautifulSoup(frag, "html.parser").get_text(strip=True):
            results.append(frag)
            continue
        # break into smaller chunks if necessary (按字符)
        max_chars = 3000
        if len(frag) <= max_chars:
            chunks = [frag]
        else:
            soup = BeautifulSoup(frag, "html.parser")
            parts = [str(p) for p in soup.find_all(['p','div','li'], recursive=False) if p.get_text(strip=True)]
            if not parts:
                parts = [frag[i:i+max_chars] for i in range(0, len(frag), max_chars)]
            chunks = parts
        translated_chunks = []
        for chunk in chunks:
            prompt = (
                f"Translate the following HTML fragment from {src_lang} to {tgt_lang}. "
                "Only translate human-readable text nodes. Preserve all HTML tags and attributes exactly. "
                "Do NOT add notes or commentary. Do NOT wrap outputs with extra explanation. Return the translated HTML fragment only.\n\n"
                f"{chunk}"
            )
            payload = {
                "model": "deepseek-chat",
                "messages": [{"role":"user","content":prompt}],
                "temperature": 0.0
            }
            try:
                resp = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, timeout=120)
                if resp.status_code == 200:
                    j = resp.json()
                    txt = j['choices'][0]['message']['content']
                    txt = clean_ai_comments(txt)
                    txt = strip_empty_tags(txt)
                    translated_chunks.append(txt)
                else:
                    print(f"⚠️ DeepSeek 返回 HTTP {resp.status_code}: {resp.text[:200]}")
                    translated_chunks.append(chunk)
            except requests.exceptions.Timeout:
                print("⚠️ DeepSeek 请求超时 (120s)，翻译失败")
                raise Exception("DeepSeek request timeout")
            except Exception as e:
                print(f"⚠️ DeepSeek 请求异常: {e}")
                raise e
        results.append("".join(translated_chunks))
    return results

def deepseek_translate_plain(texts: List[str], src_lang="zh", tgt_lang="ja") -> List[str]:
    """
    翻译纯文本（用于 shape / textbox）
    """
    if not DEEPSEEK_API_KEY:
        return [f"[MOCK]{t}" for t in texts]
    headers = {"Content-Type":"application/json", "Authorization":f"Bearer {DEEPSEEK_API_KEY}"}
    out = []
    for t in texts:
        prompt = f"Translate the following text from {src_lang} to {tgt_lang}. Do not add commentary.\n\n{t}"
        payload = {"model":"deepseek-chat","messages":[{"role":"user","content":prompt}],"temperature":0.0}
        try:
            r = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, timeout=60)
            if r.status_code == 200:
                j = r.json()
                txt = j['choices'][0]['message']['content']
                txt = clean_ai_comments(txt)
                out.append(txt)
            else:
                print(f"⚠️ DeepSeek plain 返回 HTTP {r.status_code}")
                raise Exception(f"DeepSeek HTTP {r.status_code}: {r.text}")
        except Exception as e:
            print(f"⚠️ DeepSeek plain 异常: {e}")
            raise e
    return out

# ----------------------------
# 提取 drawing/media/rels（基于 zip）
# ----------------------------
def extract_drawings_media_and_rels(docx_path: str):
    drawings = []
    media = []  # e.g. word/media/image1.png
    rels = set()
    parser = etree.XMLParser(ns_clean=True, recover=True)
    with zipfile.ZipFile(docx_path, 'r') as z:
        if 'word/document.xml' not in z.namelist():
            return drawings, media, list(rels)
        doc_xml = z.read('word/document.xml')
        root = etree.fromstring(doc_xml, parser=parser)

        # collect rels referenced by document.xml
        rels_path = 'word/_rels/document.xml.rels'
        if rels_path in z.namelist():
            rels.add(rels_path)
            rels_xml = etree.fromstring(z.read(rels_path), parser=parser)
            for rel in rels_xml.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                t = rel.get('Target')
                if t and t.startswith('media/'):
                    media.append('word/' + t)
                if t and t.startswith('drawings/'):
                    rels.add('word/' + t)
                    # drawing rels
                    dr_rels = 'word/_rels/' + t + '.rels'
                    if dr_rels in z.namelist():
                        rels.add(dr_rels)

        nodes = root.findall('.//w:drawing', namespaces=WP_NS)
        idx = 0
        for d in nodes:
            idx += 1
            draw_id = f"draw{idx}"
            drawing_xml = etree.tostring(d, encoding='utf-8').decode('utf-8')
            blip = d.find('.//a:blip', namespaces=WP_NS)
            media_fn = None
            if blip is not None:
                embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                # search in rels for matching Target
                if rels_path in z.namelist():
                    rels_xml = etree.fromstring(z.read(rels_path), parser=parser)
                    for rel in rels_xml.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                        if rel.get('Id') == embed:
                            tgt = rel.get('Target')
                            if tgt and tgt.startswith('media/'):
                                media_fn = 'word/' + tgt
                                media.append(media_fn)
                                rels.add(rels_path)
            extent_el = d.find('.//wp:extent', namespaces=WP_NS)
            extent = {}
            if extent_el is not None:
                extent['cx'] = extent_el.get('cx')
                extent['cy'] = extent_el.get('cy')
            pos = {}
            posH = d.find('.//wp:positionH', namespaces=WP_NS)
            posV = d.find('.//wp:positionV', namespaces=WP_NS)
            if posH is not None:
                pos['posH'] = etree.tostring(posH, encoding='utf-8').decode('utf-8')
            if posV is not None:
                pos['posV'] = etree.tostring(posV, encoding='utf-8').decode('utf-8')
            drawings.append({'id':draw_id,'drawing_xml':drawing_xml,'media':media_fn,'extent':extent,'pos':pos})

        # collect all media
        for name in z.namelist():
            if name.startswith('word/media/'):
                if name not in media:
                    media.append(name)

        # collect other rels
        for name in z.namelist():
            if name.startswith('word/_rels/') or '/_rels/' in name:
                rels.add(name)

    return drawings, list(dict.fromkeys(media)), list(rels)

# ----------------------------
# shapes/textbox 文本提取（低层 xml）
# ----------------------------
def extract_shapes_texts(docx_path: str):
    items = []
    parser = etree.XMLParser(ns_clean=True, recover=True)
    with zipfile.ZipFile(docx_path, 'r') as z:
        parts = [n for n in z.namelist() if n.startswith('word/') and n.endswith('.xml')]
        idx = 0
        for part in parts:
            try:
                content = z.read(part)
            except KeyError:
                continue
            try:
                root = etree.fromstring(content, parser=parser)
            except Exception:
                continue
            nodes = root.xpath('//w:drawing//w:t | //v:shape//w:t | //wps:txbx//w:t', namespaces=WP_NS)
            if nodes:
                tree = etree.ElementTree(root)
                for n in nodes:
                    txt = (n.text or "").strip()
                    if txt:
                        idx += 1
                        xpath = tree.getpath(n)
                        items.append({'id':f"shape{idx}", 'part':part, 'xpath':xpath, 'text':n.text})
    return items

# ----------------------------
# paragraph -> HTML & back
# ----------------------------
def paragraph_to_html(paragraph) -> str:
    parts = []
    for run in paragraph.runs:
        t = run.text or ""
        if not t:
            continue
        if run.bold:
            t = f"<b>{t}</b>"
        if run.italic:
            t = f"<i>{t}</i>"
        if run.underline:
            t = f"<u>{t}</u>"
        if run.font and run.font.color and run.font.color.rgb:
            t = f"<span style='color:#{run.font.color.rgb}'>{t}</span>"
        parts.append(t)
    return "<p>" + "".join(parts) + "</p>"

def paragraph_to_html_with_drawings(paragraph, drawings_slice):
    core = paragraph_to_html(paragraph)[3:-4]
    p_xml = paragraph._p
    drawing_nodes = list(p_xml.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
    for i, _ in enumerate(drawing_nodes):
        if i < len(drawings_slice):
            d = drawings_slice[i]
            style = ""
            width_px = height_px = None
            if d.get('extent',{}).get('cx'):
                try:
                    w_in = emu_to_inches(d['extent']['cx'])
                    h_in = emu_to_inches(d['extent']['cy'])
                    if w_in:
                        width_px = int(round(w_in*96))
                    if h_in:
                        height_px = int(round(h_in*96))
                except:
                    pass
            if width_px:
                style += f"width:{width_px}px;"
            if height_px:
                style += f"height:{height_px}px;"
            img_tag = f"<img data-draw-id='{d['id']}' src='{d.get('media','')}' style='{style}' data-pos='{json.dumps(d.get('pos',{}))}' />"
            core += img_tag + f"[[DRAW:{d['id']}]]"
    return f"<p>{core}</p>"

def extract_doc_to_html_structure_with_drawings(docx_path: str):
    doc = Document(docx_path)
    drawings, media_files, rels_files = extract_drawings_media_and_rels(docx_path)
    structure = []
    draw_idx = 0
    for i, para in enumerate(doc.paragraphs):
        has_draw = len(para._p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')) > 0
        if para.text.strip() or has_draw:
            slice_draws = drawings[draw_idx:draw_idx+4]
            html = paragraph_to_html_with_drawings(para, slice_draws)
            structure.append({'type':'paragraph','index':i,'text':html})
            draw_idx += len(para._p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
    # tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_html = " ".join(paragraph_to_html_with_drawings(p,[]) for p in cell.paragraphs)
                structure.append({'type':'table_cell','table_index':t_idx,'row_index':r_idx,'col_index':c_idx,'text':cell_html})
    # headers/footers
    for s_idx, sec in enumerate(doc.sections):
        for i, para in enumerate(sec.header.paragraphs):
            if para.text.strip():
                structure.append({'type':'header','section_index':s_idx,'index':i,'text':paragraph_to_html_with_drawings(para,[])})
        for i, para in enumerate(sec.footer.paragraphs):
            if para.text.strip():
                structure.append({'type':'footer','section_index':s_idx,'index':i,'text':paragraph_to_html_with_drawings(para,[])})
    return structure, drawings, media_files, rels_files

def html_to_paragraph(html: str, paragraph):
    """
    递归安全地将 HTML 写入 paragraph（尽量保留简单样式）
    """
    soup = BeautifulSoup(html, "html.parser")

    def add_run_with_style(text, para, tag_name=None, tag_obj=None):
        r = para.add_run(text)
        if tag_name == 'b':
            r.bold = True
        if tag_name == 'i':
            r.italic = True
        if tag_name == 'u':
            r.underline = True
        if tag_name == 'span' and tag_obj is not None and tag_obj.has_attr('style') and 'color' in tag_obj['style']:
            try:
                color = tag_obj['style'].split('color:')[1].split(';')[0].strip().replace('#','')
                r.font.color.rgb = RGBColor.from_string(color)
            except Exception:
                pass

    def apply_node(node, para, parent_tag_name=None, parent_tag=None):
        if isinstance(node, NavigableString):
            txt = str(node)
            if txt.strip():
                add_run_with_style(txt, para, parent_tag_name, parent_tag)
            return
        cur_tag = getattr(node, 'name', None)
        for child in node.children:
            if isinstance(child, NavigableString):
                txt = str(child)
                if txt and txt.strip():
                    add_run_with_style(txt, para, cur_tag or parent_tag_name, node if cur_tag else parent_tag)
            else:
                apply_node(child, para, cur_tag or parent_tag_name, node if cur_tag else parent_tag)

    # clear paragraph
    try:
        paragraph._p.clear()
    except Exception:
        paragraph.text = ""
    for el in soup.contents:
        apply_node(el, paragraph)

# ----------------------------
# apply shape translations to extracted dir (xml parts)
# ----------------------------
def apply_shape_translations_to_extracted_dir(extracted_dir: str, shapes_translations: List[Dict[str,Any]]):
    parser = etree.XMLParser(ns_clean=True, recover=True)
    for st in shapes_translations:
        part = st.get('part')
        xpath = st.get('xpath')
        translated = st.get('translated','')
        part_path = Path(extracted_dir) / part
        if not part_path.exists():
            print(f"⚠️ 要写回的 part 不存在: {part}")
            continue
        try:
            tree = etree.parse(str(part_path), parser=parser)
            root = tree.getroot()
            nodes = root.xpath(xpath, namespaces=WP_NS)
            if not nodes:
                # fallback：替换第一个 w:t
                all_t = root.xpath('//w:t', namespaces=WP_NS)
                if all_t:
                    all_t[0].text = translated
                    tree.write(str(part_path), xml_declaration=True, encoding='utf-8', standalone='yes')
                else:
                    print(f"⚠️ 未找到任何 w:t 来替换 part={part}")
                continue
            for n in nodes:
                n.text = translated
            tree.write(str(part_path), xml_declaration=True, encoding='utf-8', standalone='yes')
        except Exception as e:
            print(f"⚠️ 写 shape 到 xml 失败: {e}")

# ----------------------------
# Rebuild final docx: preserve media, drawings, rels; ensure table merges
# ----------------------------
def rebuild_final_docx(translated_structure: List[Dict[str,Any]], original_docx: str,
                       drawings: List[Dict[str,Any]], media_files: List[str], rels_files: List[str],
                       output_docx: str):
    # 1) build tmp docx using python-docx
    tmp_doc = Document()
    tables_map = {}
    for item in translated_structure:
        t = item.get('type')
        if t == 'paragraph':
            p = tmp_doc.add_paragraph()
            html_to_paragraph(item.get('translated',''), p)
        elif t == 'table_cell':
            t_idx = item['table_index']
            if t_idx not in tables_map:
                # create new table with same dims as original if possible
                try:
                    orig = Document(original_docx)
                    orig_tbl = orig.tables[t_idx]
                    new_tbl = tmp_doc.add_table(rows=len(orig_tbl.rows), cols=len(orig_tbl.columns))
                    new_tbl.style = 'Table Grid'
                    tables_map[t_idx] = new_tbl
                except Exception:
                    new_tbl = tmp_doc.add_table(rows=1, cols=1)
                    tables_map[t_idx] = new_tbl
            new_tbl = tables_map.get(t_idx)
            cell = new_tbl.cell(item['row_index'], item['col_index'])
            # clear cell
            cell.text = ""
            html_to_paragraph(item.get('translated',''), cell.paragraphs[0])
        elif t == 'header':
            try:
                sec = tmp_doc.sections[0]
                sec.header.add_paragraph(BeautifulSoup(item.get('translated',''),'html.parser').get_text())
            except Exception:
                tmp_doc.add_paragraph(BeautifulSoup(item.get('translated',''),'html.parser').get_text())
        elif t == 'footer':
            try:
                sec = tmp_doc.sections[0]
                sec.footer.add_paragraph(BeautifulSoup(item.get('translated',''),'html.parser').get_text())
            except Exception:
                tmp_doc.add_paragraph(BeautifulSoup(item.get('translated',''),'html.parser').get_text())
        elif t == 'shape':
            # place-holder handled in xml rewrite
            pass

    tmp_docx = Path(output_docx).with_suffix('.tmp.docx')
    tmp_doc.save(str(tmp_docx))

    # 2) unzip original and tmp into temp dirs
    tmp_dir_orig = tempfile.mkdtemp()
    tmp_dir_tmp = tempfile.mkdtemp()
    with zipfile.ZipFile(original_docx, 'r') as z:
        z.extractall(tmp_dir_orig)
    with zipfile.ZipFile(str(tmp_docx), 'r') as z:
        z.extractall(tmp_dir_tmp)

    # 3) replace drawing placeholders [[DRAW:...]] in tmp document.xml with original drawing xml
    tmp_doc_xml = Path(tmp_dir_tmp) / 'word' / 'document.xml'
    if tmp_doc_xml.exists():
        content = tmp_doc_xml.read_text(encoding='utf-8')
        for d in drawings:
            token = f"[[DRAW:{d['id']}]]"
            if token in content:
                content = content.replace(token, d['drawing_xml'])
        tmp_doc_xml.write_text(content, encoding='utf-8')

    # 4) write shape translations back into original extracted dir parts (so that rels & structure preserved)
    shape_translations = []
    for it in translated_structure:
        if it.get('type') == 'shape':
            shape_translations.append({
                'part': it.get('part'),
                'xpath': it.get('xpath'),
                'translated': BeautifulSoup(it.get('translated',''),'html.parser').get_text()
            })
    apply_shape_translations_to_extracted_dir(tmp_dir_orig, shape_translations)

    # 5) ensure table merges reflect original gridSpan/vMerge
    # We'll inspect original tables xml and apply merges on tmp_doc_tmp document.xml by editing xml
    try:
        # parse original and tmp document.xml and apply merge attributes to tmp
        parser = etree.XMLParser(ns_clean=True, recover=True)
        orig_doc_xml = etree.parse(str(Path(tmp_dir_orig) / 'word' / 'document.xml'), parser=parser)
        tmp_doc_xml_tree = etree.parse(str(Path(tmp_dir_tmp) / 'word' / 'document.xml'), parser=parser)
        # For each tbl in orig, copy tcPr gridSpan/vMerge to corresponding table cell in tmp by matching order.
        orig_tbls = orig_doc_xml.findall('.//w:tbl', namespaces=WP_NS)
        tmp_tbls = tmp_doc_xml_tree.findall('.//w:tbl', namespaces=WP_NS)
        for i, orig_tbl in enumerate(orig_tbls):
            if i >= len(tmp_tbls):
                break
            tmp_tbl = tmp_tbls[i]
            orig_tcs = orig_tbl.findall('.//w:tc', namespaces=WP_NS)
            tmp_tcs = tmp_tbl.findall('.//w:tc', namespaces=WP_NS)
            for j, o_tc in enumerate(orig_tcs):
                if j >= len(tmp_tcs):
                    break
                t_tc = tmp_tcs[j]
                o_tcPr = o_tc.find('w:tcPr', namespaces=WP_NS)
                if o_tcPr is None:
                    continue
                # copy gridSpan
                g = o_tcPr.find('w:gridSpan', namespaces=WP_NS)
                if g is not None:
                    # if tmp lacks tcPr, create
                    t_tcPr = t_tc.find('w:tcPr', namespaces=WP_NS)
                    if t_tcPr is None:
                        t_tcPr = etree.SubElement(t_tc, '{%s}tcPr' % WP_NS['w'])
                    # overwrite or set gridSpan
                    existing = t_tcPr.find('w:gridSpan', namespaces=WP_NS)
                    if existing is not None:
                        existing.set('val', g.get('val'))
                    else:
                        new = etree.SubElement(t_tcPr, '{%s}gridSpan' % WP_NS['w'])
                        new.set('val', g.get('val'))
                # copy vMerge
                v = o_tcPr.find('w:vMerge', namespaces=WP_NS)
                if v is not None:
                    t_tcPr = t_tc.find('w:tcPr', namespaces=WP_NS)
                    if t_tcPr is None:
                        t_tcPr = etree.SubElement(t_tc, '{%s}tcPr' % WP_NS['w'])
                    existing = t_tcPr.find('w:vMerge', namespaces=WP_NS)
                    if existing is not None:
                        existing.set('val', v.get('val') if v.get('val') is not None else 'continue')
                    else:
                        newv = etree.SubElement(t_tcPr, '{%s}vMerge' % WP_NS['w'])
                        if v.get('val'):
                            newv.set('val', v.get('val'))
        # write tmp_doc_xml_tree back
        tmp_doc_xml_tree.write(str(Path(tmp_dir_tmp) / 'word' / 'document.xml'), xml_declaration=True, encoding='utf-8', standalone='yes')
    except Exception as e:
        print(f"⚠️ 合并表格属性时发生异常（非致命）：{e}")

    # 6) create final zip: start with tmp_dir_tmp, then overlay original media/rels/drawings and updated parts from tmp_dir_orig
    final_path = output_docx
    with zipfile.ZipFile(final_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        # add tmp tree
        for folder, _, files in os.walk(tmp_dir_tmp):
            for fname in files:
                full = Path(folder) / fname
                arcname = str(full.relative_to(tmp_dir_tmp)).replace('\\','/')
                zout.write(str(full), arcname)
        # overlay relevant original parts (media, drawings, rels, smartArt, vml, theme)
        for folder, _, files in os.walk(tmp_dir_orig):
            for fname in files:
                full = Path(folder) / fname
                arcname = str(full.relative_to(tmp_dir_orig)).replace('\\','/')
                if arcname in zout.namelist():
                    continue
                if arcname.startswith('word/media/') or arcname.startswith('word/drawings/') \
                   or arcname.startswith('word/_rels/') or arcname.startswith('word/smartArt/') \
                   or 'vmlDrawing' in arcname or arcname.startswith('word/theme/'):
                    zout.write(str(full), arcname)
                else:
                    if arcname not in zout.namelist():
                        zout.write(str(full), arcname)

    # cleanup temp
    try:
        tmp_docx.unlink()
    except Exception:
        pass
    shutil.rmtree(tmp_dir_tmp, ignore_errors=True)
    shutil.rmtree(tmp_dir_orig, ignore_errors=True)
    print("✅ 重建完成（已尽力保留 media/drawings/rels；若文档含复杂 SmartArt/浮动，建议人工检查）")
    return final_path

# ----------------------------
# 主流程：extract -> translate -> rebuild
# ----------------------------
def translate_doc(input_path: str, output_path: str, src_lang: str = "zh", tgt_lang: str = "ja"):
    """
    input_path: 原始 docx
    output_path: 目标 docx
    src_lang/tgt_lang: 语言标识，例如 'zh'->'ja'
    """
    print(f"开始翻译：{input_path} -> {output_path} ({src_lang} -> {tgt_lang})")

    # 1) 高层提取（段落/表格）并提取 drawing/media/rels
    structure, drawings, media_files, rels_files = extract_doc_to_html_structure_with_drawings(input_path)

    # 2) 低层提取 shapes/textbox
    shapes = extract_shapes_texts(input_path)
    for s in shapes:
        structure.append({'type':'shape','id':s['id'],'part':s['part'],'xpath':s['xpath'],'text':f"<p>{s['text']}</p>"})

    # 3) 翻译 HTML 类型项
    html_items = [it for it in structure if it['type'] in ('paragraph','table_cell','header','footer')]
    html_texts = [strip_empty_tags(it['text']) for it in html_items]
    translated_htmls = deepseek_translate_html(html_texts, src_lang, tgt_lang)
    for it, tr in zip(html_items, translated_htmls):
        it['translated'] = tr

    # 4) 翻译 shape 文本（plain）
    shape_items = [it for it in structure if it['type']=='shape']
    if shape_items:
        shape_texts = [BeautifulSoup(it['text'],'html.parser').get_text() for it in shape_items]
        translated_shapes = deepseek_translate_plain(shape_texts, src_lang, tgt_lang)
        for it, tr in zip(shape_items, translated_shapes):
            it['translated'] = tr

    # 5) Rebuild final docx
    final = rebuild_final_docx(structure, input_path, drawings, media_files, rels_files, output_path)

    # 6) 提示
    if any(d.get('pos') for d in drawings):
        print("⚠️ 检测到浮动/锚点图片；程序尽力恢复 drawing XML 和 media，但位置可能仍需人工微调。")
    with zipfile.ZipFile(input_path, 'r') as z:
        if any(n.startswith('word/smartArt/') for n in z.namelist()):
            print("⚠️ 检测到 SmartArt：复杂样式可能无法完全还原。若需 1:1 保真，建议使用 Aspose 或 Office 专业工具。")

    print(f"完成：{final}")
    return final

# ----------------------------
# CLI
# ----------------------------
if __name__ == "__main__":
    import sys
    if len(sys.argv) < 3:
        print("用法: python translator_html.py input.docx output.docx [src] [tgt]")
        sys.exit(1)
    inp = sys.argv[1]
    outp = sys.argv[2]
    src = sys.argv[3] if len(sys.argv) > 3 else 'zh'
    tgt = sys.argv[4] if len(sys.argv) > 4 else 'ja'
    translate_doc(inp, outp, src, tgt)
